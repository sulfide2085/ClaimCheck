from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


SOURCE = Path("docs/ClaimCheck-plan-revised.md")
OUTPUT = Path("out/ClaimCheck-跨境电商竞品卖点验证助手-改进版.docx")


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.size = Pt(11)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.size = Pt(11)


def build_doc() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("ClaimCheck：跨境电商竞品卖点验证助手（改进版）")
    title_run.bold = True
    title_run.font.size = Pt(16)

    for raw in lines[1:]:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(13)
            continue
        if line.startswith("### "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(12)
            continue
        if line.startswith("- "):
            add_bullet(document, line[2:])
            continue
        add_paragraph(document, line)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_doc()
